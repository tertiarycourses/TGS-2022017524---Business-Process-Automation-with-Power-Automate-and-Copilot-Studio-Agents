#!/usr/bin/env python3
"""
WSQ Assessments (DOCX) for
"Business Process Automation with Power Automate and Copilot Studio Agents"
(TGS-2022017524).

Two instruments, each as a question paper + a model-answer / marking guide:

  1. Written Assessment (WA / SAQ) — tests KNOWLEDGE. 6 open-ended
     short-answer questions (K1-K6) drawn from Modules 1-3 and the slides.
     Duration: 1 hour, open book.
  2. Practical Performance (PP) Assessment — tests PRACTICAL ability.
     One ACME Pte Ltd scenario with 3 tasks (criteria A1-A7) whose model
     answers are the hands-on lab build steps (Labs 1-15).
     Duration: 1 hour, open book.

The question counts (WA 6 / PP 3) intentionally match the previous versions.
Writes four files into ../assessemnt/.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# script lives at .claude/skills/wsq-assessment/ — repo root is 3 levels up
REPO = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".."))
OUT = os.path.join(REPO, "assessemnt")
TITLE = "Business Process Automation with Power Automate and Copilot Studio Agents"
COURSE_CODE = "TGS-2022017524"
ORG = "Tertiary Infotech Academy Pte Ltd"
UEN = "201200696W"
DARK = RGBColor(0x16, 0x1B, 0x26); BRAND = RGBColor(0x1F, 0x6F, 0xEB); GREY = RGBColor(0x55, 0x5B, 0x66)

# ================================================================ doc helpers
def new_doc():
    d = Document(); n = d.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11); return d

def line(d, text="", bold=False, size=11, color=DARK, after=6, align=None):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p

def runs(d, segments, after=6, style=None):
    p = d.add_paragraph(style=style); p.paragraph_format.space_after = Pt(after)
    for text, bold in segments:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(11); r.font.name = "Arial"; r.font.color.rgb = DARK
    return p

def bullet(d, text):
    return runs(d, [(text, False)], after=3, style="List Bullet")

def numbered(d, text):
    return runs(d, [(text, False)], after=3, style="List Number")

def cover(d, subtitle):
    for _ in range(4): d.add_paragraph()
    line(d, ORG, bold=True, size=13, align=AL.CENTER, after=2)
    line(d, f"UEN: {UEN}", size=10, color=GREY, align=AL.CENTER, after=20)
    line(d, subtitle.upper(), bold=True, size=24, color=BRAND, align=AL.CENTER, after=12)
    line(d, "For", size=12, color=GREY, align=AL.CENTER, after=8)
    line(d, TITLE, bold=True, size=18, align=AL.CENTER, after=14)
    line(d, f"TGS Ref No: {COURSE_CODE}", size=12, color=GREY, align=AL.CENTER, after=20)
    line(d, "Conducted by", size=12, color=GREY, align=AL.CENTER, after=6)
    line(d, ORG, bold=True, size=13, align=AL.CENTER, after=2)
    line(d, f"UEN: {UEN}", size=10, color=GREY, align=AL.CENTER, after=2)
    d.add_page_break()

def title_block(d, subtitle):
    line(d, TITLE, bold=True, size=15, color=BRAND, after=2, align=AL.CENTER)
    line(d, subtitle, bold=True, size=12, color=DARK, after=2, align=AL.CENTER)
    line(d, f"Course Code: {COURSE_CODE}", size=10, color=GREY, after=12, align=AL.CENTER)

# Page 2 blocks (house standard — see the wsq-assessment skill's build_assessment.py):
# Trainee Information + Instructions to Candidate + Grading, then a page break so the
# questions/tasks always begin on page 3.
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

def trainee_info(d):
    line(d, "Trainee Information", bold=True, size=12, after=6)
    for t in ["Trainee Name (as per NRIC): ______________________________________",
              "Last 3 digits and alphabet of NRIC / FIN: ____________________",
              "Date: ____________________"]:
        p = line(d, t, after=6); p.paragraph_format.line_spacing = 2.0

def instructions_block(d, time_text, extra):
    line(d, "Instructions to Candidate", bold=True, size=12, after=6)
    items = ["This is an individual exercise.",
             "This is an open-book assessment (slides and Learner Guide may be used).",
             f"A total of {time_text} is given to complete this assessment.",
             extra,
             "Complete your answers on the document provided and upload the completed "
             "answers to the LMS at https://lms-tms.tertiaryinfotech.com/."] + BRIEFING
    for i, s in enumerate(items, 1):
        line(d, f"{i}.  {s}", after=3)
    line(d, "", after=2)

def grading_block(d, statement):
    line(d, "Grading", bold=True, size=12, after=6)
    line(d, statement, after=10)
    for t in ["Grade: _______  (C / NYC)",
              "Assessor Name: __________________________   Assessor NRIC: ________________",
              "Date: ________________________                    Signature: ____________________"]:
        p = line(d, t, after=6); p.paragraph_format.line_spacing = 2.0

def answer_box(d, height_cm=5.0, hint=""):
    t = d.add_table(rows=1, cols=1); t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.text = ""
    if hint:
        p = cell.paragraphs[0]; r = p.add_run(hint); r.font.size = Pt(9); r.font.color.rgb = GREY; r.font.name = "Arial"
    trPr = t.rows[0]._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight"); h.set(qn("w:val"), str(int(height_cm * 567))); h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)
    d.add_paragraph()

# ================================================================ WRITTEN (6 questions, K1-K6)
# (tag, question, model answer points, source)
WRITTEN = [
    ("K1",
     "ACME Pte Ltd still records every customer enquiry by hand: a customer emails, someone copies the details "
     "into a spreadsheet, then a salesperson is told to reply. Explain what business workflow automation is, and "
     "describe how the Trigger → Actions → Output model would apply if ACME automated this enquiry process.",
     ["Business workflow automation = using software to run a repeatable series of business steps automatically, "
      "instead of a person doing each step by hand (Module 1).",
      "Trigger — the event that starts the workflow: a new enquiry email arrives in the sales mailbox.",
      "Actions — what the workflow does: extract the details, add a row into the Excel enquiries table, "
      "notify/email the salesperson.",
      "Output — the result produced/passed on: the logged row, the notification, and the customer reply.",
      "(Any equivalent mapping of trigger/actions/output to the enquiry process is acceptable.)"],
     "Module 1 · Slides 16–27"),
    ("K2",
     "ACME's operations manager asks you which of the company's processes should be automated first. State three "
     "traits that make a task a good candidate for automation, and give one example of a business process that "
     "shows all three traits.",
     ["Repetitive — the task is done the same way many times (e.g. daily/weekly).",
      "Rule-based — it follows clear \"if this, then that\" logic with no judgement calls.",
      "Time-consuming — it involves manual copying, chasing or re-typing.",
      "Example: logging every incoming invoice to a spreadsheet and chasing the manager for approval "
      "(or any similar process that is repetitive, rule-based and time-consuming)."],
     "Module 1 · Slides 17–19"),
    ("K3",
     "Describe THREE types of Power Automate cloud flow covered in the course, the event that starts each one, and "
     "give a business example of when you would use each type.",
     ["Instant / manual flow — started when a person presses Run; e.g. a one-click email or test workflow (Lab 1).",
      "Scheduled flow — started by a Recurrence timetable; e.g. a Monday 9:00 AM weekly reminder or stock check (Lab 4).",
      "Automated flow — started by an event: a new form response, a new email arriving, or a file being uploaded; "
      "e.g. logging each form submission to Excel (Lab 5).",
      "(Agent flow — started when a Copilot Studio agent calls the flow as a tool — is also acceptable as one of the three.)"],
     "Module 2 · Slides 28–40"),
    ("K4",
     "Your colleague's flow fails at the Send an email step with an \"Unauthorized\" error. Explain what a connector "
     "is in Power Automate, name three connectors used in this course, and state the likely cause of the error and "
     "how to fix it.",
     ["A connector links a flow to a service and provides its triggers and actions; a connection is your signed-in "
      "account for that connector.",
      "Any three of: Office 365 Outlook, Excel Online (Business), Approvals, Microsoft Forms, OneDrive for Business, "
      "Microsoft Teams.",
      "Cause: the connection is broken/expired, or the signed-in account has no mailbox / lacks rights.",
      "Fix: reconnect the Office 365 Outlook connection with a valid, mailbox-enabled tenant account "
      "(green tick = ready, red = reconnect)."],
     "Module 2 · Slides 33–40"),
    ("K5",
     "ACME wants a customer-service agent built in Microsoft Copilot Studio. Name the main building blocks of a "
     "Copilot Studio agent and what each is for, and explain how Knowledge (RAG / grounding) keeps the agent's "
     "answers accurate.",
     ["Instructions — plain-language directions that shape the agent's behaviour and tone.",
      "Knowledge — documents/websites the agent answers from.",
      "Topics — conversation flows (Ask a question, Condition, Message nodes) for structured dialogues.",
      "Tools — things the agent can do, including Power Automate (agent) flows.",
      "Variables — where captured answers are stored and passed onward.",
      "RAG (Retrieval-Augmented Generation): the agent retrieves relevant passages from the uploaded sources and "
      "generates an answer grounded in them — preventing made-up answers (hallucination)."],
     "Module 3 · Slides 50–59 · Labs 6–7"),
    ("K6",
     "ACME's Procurement Assistant agent must hand a purchase request to a Power Automate flow that logs the request "
     "and emails the manager. Name the trigger and the closing action the flow must use, and describe how data is "
     "passed between the agent and the flow. State one best practice for this agent + flow integration.",
     ["Trigger: \"When an agent calls the flow\"; closing action: \"Respond to the agent\".",
      "The agent's topic adds the flow as a tool; the agent's variables (e.g. item, quantity, cost) are mapped to the "
      "flow's input parameters, and the flow's outputs are returned to the agent via Respond to the agent.",
      "Both the agent and the flow must live in the SAME environment or the agent cannot see the flow.",
      "Best practice (any one): keep tool names/descriptions clear so orchestration picks the right tool; validate "
      "inputs with Ask a question (Number entity); notify on every branch so there are no dead ends; test the happy "
      "path and every branch."],
     "Modules 3–4 · Slides 60–64, 75–81 · Labs 10, 14"),
]

WA_NAME = f"WA (SAQ) - {TITLE}"
PP_NAME = f"PP Assessment - {TITLE}"

def build_written_paper():
    d = new_doc()
    cover(d, "Written Assessment (SAQ)")
    title_block(d, "Written Assessment (SAQ) — Knowledge")
    # Page 2: trainee info + instructions + grading; questions begin on page 3.
    trainee_info(d)
    instructions_block(d, "1 hour", "Answer ALL 6 questions in the boxes provided.")
    grading_block(d, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
    d.add_page_break()
    line(d, "Short-Answer Questions (Knowledge)", bold=True, size=12, after=8)
    for i, (tag, q, _pts, _src) in enumerate(WRITTEN, 1):
        runs(d, [(f"Question {i}: ", True), (q, False), (f"  ({tag})", True)], after=6)
        answer_box(d, 5.0)
    out = os.path.join(OUT, f"{WA_NAME}.docx"); d.save(out); return out

def build_written_answers():
    d = new_doc()
    cover(d, "Answers to Written Assessment (SAQ)")
    title_block(d, "Written Assessment (SAQ) — Model Answers / Marking Guide")
    line(d, "Assessor note: answers are model points — accept equivalent wording. A response is Competent for a "
            "question when it covers the substance of the points shown.", color=GREY, after=10)
    for i, (tag, q, pts, src) in enumerate(WRITTEN, 1):
        runs(d, [(f"Question {i} ({tag}): ", True), (q, False)], after=4)
        for pt in pts: bullet(d, pt)
        runs(d, [("Source: ", True), (src, False)], after=12)
    out = os.path.join(OUT, f"Answers to {WA_NAME}.docx"); d.save(out); return out

# ================================================================ PRACTICAL (3 tasks, A1-A7)
SCENARIO = (
    "Scenario: ACME Pte Ltd is a growing trading company. Staff currently email purchase requests to the "
    "operations manager, who copies them into a spreadsheet and replies by hand. ACME wants you — as taught in "
    "this course — to automate the process end to end: a Power Automate flow to log and approve requests, a "
    "Copilot Studio agent to capture requests in a conversation, and the agent connected to the flow. "
    "Build everything in your Course Sandbox environment (Power Automate and Copilot Studio must use the SAME "
    "environment)."
)

PP_TASKS = [
    ("Task 1 — Build the approval & logging flow (Power Automate)",
     "In make.powerautomate.com, create an instant cloud flow named PP Purchase Approval that: "
     "(a) is started manually with two text inputs — Item and Cost; "
     "(b) adds a row into an Excel Online (Business) table (columns: Item, Cost, Date — use utcNow() via the fx "
     "expression editor for Date); "
     "(c) starts an approval with \"Start and wait for an approval\" assigned to yourself; and "
     "(d) uses a Condition on the approval Outcome to send an \"Approved\" email if Approve, otherwise a "
     "\"Rejected\" email. Run the flow once and approve it.",
     "Take a screenshot of the whole flow in the designer AND a screenshot of the successful run, and paste them "
     "in the box below: (A1, A2, A3)"),
    ("Task 2 — Build the Procurement Assistant agent (Copilot Studio)",
     "In copilotstudio.microsoft.com (same environment), create an agent named PP Procurement Assistant with: "
     "(a) instructions that make it a polite internal procurement helper; "
     "(b) a topic New Request that uses Ask a question nodes to capture the Item (text) and the Cost (Number "
     "entity) into variables; and "
     "(c) a Message node that confirms the captured values back to the user. "
     "Test the conversation in the Test pane.",
     "Take a screenshot of the topic (showing the question nodes and variables) AND the Test pane conversation, "
     "and paste them in the box below: (A4, A5)"),
    ("Task 3 — Connect the agent to the flow (end to end)",
     "Make the agent do the work: (a) create an agent flow with the trigger \"When an agent calls the flow\" that "
     "receives Item and Cost as inputs, logs them to your Excel table, sends the notification email, and returns a "
     "confirmation with \"Respond to the agent\"; (b) add this flow as a tool at the end of your New Request topic "
     "and map the topic's variables to the flow inputs; and (c) test end to end in the Test pane — the row must "
     "appear in Excel and the email must arrive.",
     "Take a screenshot of the topic showing the tool node with the mapped inputs AND a screenshot of the Excel "
     "row / email produced by your test, and paste them in the box below: (A6, A7)"),
]

PP_ANSWERS = [
    ("Task 1 (A1, A2, A3) — model build steps (mirrors Labs 1, 2, 3)", [
        "Create → Instant cloud flow → name PP Purchase Approval → trigger \"Manually trigger a flow\" (Lab 1).",
        "In the trigger, Add an input → Text twice: Item, Cost (Lab 3 pattern).",
        "+ Add an action → \"Add a row into a table\" (Excel Online (Business)) → pick Location / Document Library / "
        "File / Table; map Item and Cost tokens; for Date use the fx expression editor with "
        "formatDateTime(utcNow(),'yyyy-MM-dd HH:mm') so it becomes a token, not literal text (Lab 2).",
        "+ Add an action → \"Start and wait for an approval\" → Approval type: Approve/Reject – First to respond; "
        "Assigned to: yourself picked from the people picker (must be a tenant user) (Lab 3).",
        "+ Add an action → Condition → Outcome is equal to Approve. True: Send an email (V2) \"Approved\". "
        "False: Send an email (V2) \"Rejected\" (Lab 3).",
        "Save → Test → Manually → enter Item/Cost → approve from the Approvals hub → run shows all green ticks.",
        "Competent when: flow runs successfully; row logged with a real date; approval received and actioned; "
        "correct email on each branch."]),
    ("Task 2 (A4, A5) — model build steps (mirrors Labs 6, 9)", [
        "Copilot Studio → confirm the environment selector (top-right) shows Course Sandbox — the same environment "
        "as Power Automate (Labs 0, 6).",
        "Left nav Agents → Create blank agent → the agent's Overview page opens; Details → Edit → name it "
        "PP Procurement Assistant; Instructions → Edit → \"You are ACME's polite internal procurement helper…\" "
        "→ Save (Lab 6).",
        "Topics → + Add a topic → From blank → name New Request; give it a clear description (generative "
        "orchestration selects topics by description) (Lab 9).",
        "Ask a question node 1: \"What item do you need?\" → Identify: User's entire response → save to variable "
        "item (Lab 9).",
        "Ask a question node 2: \"What is the estimated cost?\" → Identify: Number → save to variable cost — the "
        "Number entity re-asks on invalid input (Lab 9).",
        "Message node: \"You requested {item} at ${cost}. Thank you!\" using the variable chips.",
        "Test in the Test pane: trigger the topic, answer both questions, see the confirmation message.",
        "Competent when: agent exists with sensible instructions; both variables captured with the right entity "
        "types; confirmation echoes the values."]),
    ("Task 3 (A6, A7) — model build steps (mirrors Labs 8, 10)", [
        "In the New Request topic, after the questions: + → Add a tool → New agent flow (opens Power Automate with "
        "the \"When an agent calls the flow\" trigger) (Labs 10, 14).",
        "In the trigger, Add an input → Text: Item; Number: Cost.",
        "Add \"Add a row into a table\" (map the trigger tokens) and \"Send an email (V2)\" (to yourself, subject "
        "\"New purchase request\", body with the tokens).",
        "End with \"Respond to the agent\" → output confirmation (e.g. \"Logged\"). Click Publish — publishing "
        "saves and publishes the agent flow in one step (Lab 10).",
        "Back in the topic, the flow appears as a tool node — map topic variables item → Item and cost → Cost "
        "(refresh if the flow is not listed; agent and flow must be in the same environment) (Lab 10).",
        "Optionally end the topic with a Message node using the flow's output.",
        "Test pane: run the whole conversation → confirm the Excel row appears and the email arrives.",
        "Competent when: agent flow uses the correct trigger/response actions; inputs mapped from the topic "
        "variables; end-to-end test produces the row and the email."]),
]

def build_pp_paper():
    d = new_doc()
    cover(d, "Practical Performance (PP) Assessment")
    title_block(d, "Practical Performance (PP) Assessment")
    # Page 2: trainee info + instructions + grading; the practical problem begins on page 3.
    trainee_info(d)
    instructions_block(d, "1 hour", "Complete ALL 3 tasks and paste the required screenshots in the boxes provided.")
    grading_block(d, "Candidate has successfully completed all the tasks for PP and is able to explain "
                     "the overall functions and features used to achieve these tasks.")
    d.add_page_break()
    line(d, "Practical Performance", bold=True, size=12, after=6)
    line(d, SCENARIO, after=10)
    for head, body, evidence in PP_TASKS:
        runs(d, [(head + ": ", True), (body, False)], after=6)
        line(d, evidence, bold=True, after=4)
        answer_box(d, 6.0, hint="Paste screenshot(s) here")
    out = os.path.join(OUT, f"{PP_NAME}.docx"); d.save(out); return out

def build_pp_answers():
    d = new_doc()
    cover(d, "Answers to Practical Performance (PP) Assessment")
    title_block(d, "Practical Performance (PP) — Model Answers / Marking Guide")
    line(d, "Assessor note: the model answers below are the in-class lab build steps. Accept any working build that "
            "meets the Competent-when criteria; the cited labs show the exact click paths.", color=GREY, after=8)
    line(d, SCENARIO, after=10)
    for head, steps in PP_ANSWERS:
        line(d, head, bold=True, after=4)
        for s in steps: numbered(d, s)
        line(d, "", after=6)
    out = os.path.join(OUT, f"Answer to {PP_NAME}.docx"); d.save(out); return out

if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    for fn in (build_written_paper, build_written_answers, build_pp_paper, build_pp_answers):
        print("Wrote", fn())
