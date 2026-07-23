#!/usr/bin/env python3
"""
Single-source Learner Guide generator for
"Microsoft Copilot Studio & Power Automate for Business Workflow Automation".

Compiles the actual lab + module markdown under ./labs/ into BOTH:
  - LEARNER-GUIDE.md                                            (repo root)
  - courseware/Microsoft Copilot Studio & Power Automate ... Learner Guide.docx
so the two are always aligned. Re-run after editing any lab.
"""
import os, re, sys

SKILL = "/Users/alfredang/.claude/skills/tertiary-learner-guide"
sys.path.insert(0, SKILL)
import prodoc
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# script lives at .claude/skills/wsq-learner-guide/ — repo root is 3 levels up
REPO = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".."))
TITLE = "Business Process Automation with Power Automate and Copilot Studio Agents"
VERSION = "3.0"
COURSE_CODE = "TGS-2022017524"
ORG = "Tertiary Infotech Academy Pte Ltd"
UEN = "201200696W"
VERSIONS = [
    ["1.0", "24 Jun 2026", "Initial release — full 3-day, 17-lab learner guide.",
     "Course Development Team"],
    ["2.0", "2 Jul 2026", "WSQ revision — new course title, labs updated to the current "
     "Copilot Studio / Power Automate UI, Course Sandbox environment, WSQ cover page.",
     "Course Development Team"],
    ["3.0", "3 Jul 2026", "Course restructured from 3 days to 2 days — Day 1: Power Automate "
     "(Labs 0-5), Day 2: Copilot Studio agents (Labs 6-11) ending with the WSQ assessment. "
     "Modules 4-5 and Labs 12-16 retired.",
     "Course Development Team"],
]

# ---- neutral (un-branded) cover + footer: this is client training, no Tertiary branding ----
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x16, 0x1B, 0x26)
def _cline(doc, text, size, bold=False, color=DARK, before=0, after=4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p
def add_cover_neutral(doc, kind, title, version, course_code):
    for _ in range(3): doc.add_paragraph()
    _cline(doc, ORG, 13, bold=True, after=2)
    _cline(doc, f"UEN: {UEN}", 10, color=GREY, after=20)
    _cline(doc, kind.upper(), 26, bold=True, color=BRAND, after=12)
    _cline(doc, "For", 12, color=GREY, after=8)
    _cline(doc, title, 20, bold=True, color=DARK, after=14)
    _cline(doc, f"TGS Ref No: {course_code}", 12, color=GREY, after=20)
    _cline(doc, "Conducted by", 12, color=GREY, after=6)
    _cline(doc, ORG, 13, bold=True, after=2)
    _cline(doc, f"UEN: {UEN}", 10, color=GREY, after=18)
    _cline(doc, f"Version {version}", 12, bold=True, color=BRAND)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def add_footer_neutral(doc, title):
    sec = doc.sections[0]; footer = sec.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.text = ""
    r = p.add_run("Page "); r.font.size = Pt(9); r.font.color.rgb = GREY
    prodoc._field(p, "PAGE", "1").font.size = Pt(9)
    r2 = p.add_run(" of "); r2.font.size = Pt(9); r2.font.color.rgb = GREY
    prodoc._field(p, "NUMPAGES", "1").font.size = Pt(9)
    sp = footer.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(f"{title}  ·  © {ORG}"); sr.font.size = Pt(7.5); sr.font.color.rgb = GREY

# ---- course structure: ordered (day-heading, [files]) ----
DAYS = [
    ("Day 1 — Foundations & Power Automate", [
        "labs/Day 1/Module 1 - Workflow Automation Concepts.md",
        "labs/Day 1/Module 2 - Introduction to Power Automate.md",
        "labs/Day 1/Lab 0 - Environment Setup/index.md",
        "labs/Day 1/Lab 1 - Automated Email Workflow/index.md",
        "labs/Day 1/Lab 2 - Excel Data Logging Workflow/index.md",
        "labs/Day 1/Lab 3 - Simple Approval Workflow/index.md",
        "labs/Day 1/Lab 4 - Scheduled Trigger Workflow/index.md",
        "labs/Day 1/Lab 5 - Form Submission Workflow/index.md",
    ]),
    ("Day 2 — Building Business Agents with Copilot Studio", [
        "labs/Day 2/Module 3 - Business Agents Concepts.md",
        "labs/Day 2/Lab 6 - Create Your First Agent/index.md",
        "labs/Day 2/Lab 7 - Add Knowledge to Your Agent/index.md",
        "labs/Day 2/Lab 8 - Add Tools and Actions/index.md",
        "labs/Day 2/Lab 9 - Sales Enquiry Assistant/index.md",
        "labs/Day 2/Lab 10 - Procurement Request Workflow/index.md",
        "labs/Day 2/Lab 11 - Automated Response Generation/index.md",
    ]),
]

# ============================================================================
# Markdown -> generic blocks
# ============================================================================
LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
def clean(t):
    return LINK.sub(r"\1", t).rstrip()

def md_to_blocks(text):
    lines = text.split("\n")
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()
        # code fence
        if s.startswith("```"):
            buf = []; i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            blocks.append(("code", "\n".join(buf))); continue
        # blank
        if not s:
            i += 1; continue
        # heading
        m = re.match(r"(#{1,6})\s+(.*)", s)
        if m:
            blocks.append(("h", len(m.group(1)), clean(m.group(2)))); i += 1; continue
        # horizontal rule
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", s):
            blocks.append(("rule",)); i += 1; continue
        # table
        if s.startswith("|") and "|" in s[1:]:
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^:?-{2,}:?$", cells[0]):  # skip separator row
                    rows.append([clean(c) for c in cells])
                i += 1
            if rows: blocks.append(("table", rows))
            continue
        # blockquote (collect consecutive)
        if s.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]).rstrip()); i += 1
            content = clean(" ".join(x for x in buf if x.strip()))
            if content: blocks.append(("note", content))
            continue
        # numbered list (with indented sub-lines merged into the step)
        if re.match(r"^\d+\.\s+", s):
            items = []
            while i < n:
                raw = lines[i]
                mm = re.match(r"^\s*\d+\.\s+(.*)", raw)
                if mm:
                    items.append(clean(mm.group(1)))
                elif re.match(r"^\s{2,}\S", raw) and not raw.strip().startswith("```"):
                    # continuation / sub-bullet -> append to current step
                    sub = re.sub(r"^\s*[-*]\s+", "", raw.strip())
                    if items: items[-1] = items[-1] + "  —  " + clean(sub)
                elif raw.strip() == "":
                    # blank: peek if list continues
                    j = i + 1
                    if j < n and re.match(r"^\s*\d+\.\s+", lines[j]):
                        i += 1; continue
                    else:
                        break
                else:
                    break
                i += 1
            blocks.append(("steps", items)); continue
        # bullet list
        if re.match(r"^[-*]\s+", s):
            items = []
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(clean(re.sub(r"^\s*[-*]\s+", "", lines[i].strip()))); i += 1
            blocks.append(("bullets", items)); continue
        # plain paragraph
        blocks.append(("p", clean(s))); i += 1
    return blocks

# ============================================================================
# Assemble the DSL block list B
# ============================================================================
B = []
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def note(t): B.append(("note", t))
def rule(): B.append(("rule",))

# Title + intro
B.append(("h1", "Learner Guide"))
p(f"Welcome! This Learner Guide takes you **click-by-click** through every hands-on lab in the WSQ course "
  f"**{TITLE}** (Course Code: {COURSE_CODE}). Over two days you go from your first Power Automate flow "
  f"to AI business agents in Microsoft Copilot Studio — and finish by connecting an agent to your flows in a "
  f"complete end-to-end automated workflow.")
p("Work through the labs **in order**: each one builds on the skills of the lab before it. Whenever you see a "
  "**Checkpoint**, stop and confirm your flow or agent behaves as described before moving on. The "
  "**Common Errors & Quick Fixes** and per-lab **Troubleshooting** tables will get you unstuck fast.")
note("Course flow at a glance — Day 1: Workflow automation concepts + Power Automate (Labs 0-5). "
     "Day 2: Business agents in Copilot Studio + agent-and-flow end-to-end workflows (Labs 6-11), "
     "then the WSQ assessment (4:00-6:00 PM).")
rule()

# Common errors reference (mirrors the deck's quick-fix slide)
h2("Common Errors & Quick Fixes")
p("Keep this handy — these are the issues learners hit most often, with the one-line fix:")
B.append(("table", [
    ["Symptom", "Cause", "Fix"],
    ["“Unauthorized” when sending email", "Outlook connection expired or the account has no mailbox",
     "Reconnect the Office 365 Outlook connection with a mailbox-enabled account; both connections must be green ✓"],
    ["Approval fails: “valid users in the organization”", "Approver typed as an external email, not a tenant user",
     "Pick the approver from the people-picker dropdown (a real user in your tenant; yourself is fine for testing)"],
    ["Date logs as literal text ‘utcNow()’", "Expression typed into the field as text",
     "Enter it via the fx / Expression editor so it becomes a coloured token"],
    ["Excel cell shows ########", "The column is only too narrow", "Auto-fit the column — the value is fine"],
    ["An unwanted ‘For each’ wraps your action", "You inserted a list/array value into a single-value field",
     "Use single-value fields (Outcome, trigger inputs); delete the For each and re-add a plain action"],
    ["Agent can’t see its flow", "Agent and flow are in different environments",
     "Set both Copilot Studio and Power Automate to the same environment (Course Sandbox)"],
]))
rule()

for day_title, files in DAYS:
    h2(day_title)
    for rel in files:
        path = os.path.join(REPO, rel)
        if not os.path.exists(path):
            print("  [missing]", rel); continue
        raw = open(path, encoding="utf-8").read()
        gblocks = md_to_blocks(raw)
        first_h1_done = False
        for gb in gblocks:
            if gb[0] == "h":
                lvl, txt = gb[1], gb[2]
                if lvl == 1 and not first_h1_done:
                    h3(txt); first_h1_done = True            # lab/module title -> Heading 2
                else:
                    p(f"**{txt}**")                           # sections / steps -> bold line
            else:
                B.append(gb)
        rule()

# ============================================================================
# Renderers (markdown + docx)
# ============================================================================
def _anchor(txt):
    # GitHub slug rules: lowercase, strip punctuation, EACH space becomes a
    # hyphen (so "A — B" and "A & B" produce double hyphens, not one).
    return re.sub(r"[^\w\- ]", "", txt.lower()).replace(" ", "-")

def _toc(blocks):
    out = ["## Table of Contents", ""]
    for b in blocks:
        if b[0] == "h2":
            out.append(f"- [{b[1]}](#{_anchor(b[1])})")
        elif b[0] == "h3":
            out.append(f"  - [{b[1]}](#{_anchor(b[1])})")
    out.append("")
    return "\n".join(out)

def render_markdown(blocks):
    out = []; injected = False
    for b in blocks:
        k = b[0]
        if k == "h1":
            out.append(f"# {b[1]}\n")
            if not injected:
                out.append(f"**Course Code:** {COURSE_CODE}  ·  **Version {VERSION}**\n")
                out.append("### Document Version Control Record\n")
                out.append("| Version | Effective Date | Summary of Changes | Author |")
                out.append("| --- | --- | --- | --- |")
                for v in VERSIONS:
                    out.append(f"| {v[0]} | {v[1]} | {v[2]} | {v[3]} |")
                out.append("")
                out.append(_toc(blocks)); injected = True
            continue
        if k == "h2": out.append(f"## {b[1]}\n")
        elif k == "h3": out.append(f"### {b[1]}\n")
        elif k == "p": out.append(f"{b[1]}\n")
        elif k == "steps": out.append("\n".join(f"{i}. {s}" for i, s in enumerate(b[1], 1)) + "\n")
        elif k == "bullets": out.append("\n".join(f"- {s}" for s in b[1]) + "\n")
        elif k == "code": out.append("```\n" + b[1] + "\n```\n")
        elif k == "table":
            rows = b[1]
            out.append("| " + " | ".join(rows[0]) + " |")
            out.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for r in rows[1:]:
                out.append("| " + " | ".join(r) + " |")
            out.append("")
        elif k == "note": out.append(f"> {b[1]}\n")
        elif k == "rule": out.append("---\n")
    return "\n".join(out).strip() + "\n"

GREY = RGBColor(0x55,0x5B,0x66)
def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); tcPr.append(shd)
def _shade_para(pr, hexc="F3F5F8"):
    ppr = pr._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); ppr.append(shd)
def _runs(par, text):
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = RGBColor(0xC7,0x25,0x4E)
        else:
            par.add_run(part)

def render_docx(blocks):
    doc = Document()
    nrm = doc.styles["Normal"]; nrm.font.name = "Arial"; nrm.font.size = Pt(11)
    prodoc.style_headings(doc)
    add_cover_neutral(doc, "Learner Guide", TITLE, VERSION, COURSE_CODE)
    prodoc.add_version_control(doc, VERSIONS)
    prodoc.add_toc(doc, levels="1-2")
    for b in blocks:
        k = b[0]
        if k == "h1": continue
        elif k == "h2":
            doc.add_paragraph(style="Heading 1").add_run(b[1])
        elif k == "h3":
            doc.add_paragraph(style="Heading 2").add_run(b[1])
        elif k == "p":
            _runs(doc.add_paragraph(), b[1])
        elif k == "steps":
            for s in b[1]:
                _runs(doc.add_paragraph(style="List Number"), s)
        elif k == "bullets":
            for s in b[1]:
                _runs(doc.add_paragraph(style="List Bullet"), s)
        elif k == "code":
            pr = doc.add_paragraph(); _shade_para(pr)
            r = pr.add_run(b[1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif k == "table":
            rows = b[1]
            t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci, val in enumerate(row):
                    cells[ci].text = ""; pp = cells[ci].paragraphs[0]
                    if ri == 0:
                        rr = pp.add_run(val); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); rr.font.size = Pt(9.5)
                        _shade(cells[ci], "1F6FEB")
                    else:
                        _runs(pp, val)
                        for rn in pp.runs: rn.font.size = Pt(9.5)
        elif k == "note":
            pr = doc.add_paragraph(); _shade_para(pr, "FFF4E5")
            rr = pr.add_run("Note:  "); rr.bold = True; rr.font.color.rgb = RGBColor(0xB5,0x6A,0x00)
            _runs(pr, b[1])
        elif k == "rule":
            pr = doc.add_paragraph(); pr.paragraph_format.space_before = Pt(2); pr.paragraph_format.space_after = Pt(2)
            ppr = pr._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"D0D7DE")
            bdr.append(bot); ppr.append(bdr)
    add_footer_neutral(doc, TITLE)
    prodoc.enable_update_fields(doc)
    return doc

md = render_markdown(B)
with open(os.path.join(REPO, "LEARNER-GUIDE.md"), "w", encoding="utf-8") as f:
    f.write(md)
out_docx = os.path.join(REPO, f"courseware/LG-{TITLE}.docx")
render_docx(B).save(out_docx)
print("Wrote LEARNER-GUIDE.md (%d blocks, %d labs/modules headings)" %
      (len(B), sum(1 for b in B if b[0]=="h3")))
print("Wrote", out_docx)
