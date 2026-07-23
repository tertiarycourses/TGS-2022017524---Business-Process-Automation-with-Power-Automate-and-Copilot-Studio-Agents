#!/usr/bin/env python3
"""
2-day WSQ Lesson Plan (DOCX) for
"Business Process Automation with Power Automate and Copilot Studio Agents"
(TGS-2022017524), Tertiary Infotech Academy Pte Ltd house format.

Daily window 9:00am - 6:00pm (1-hour lunch; tea breaks within).
Day 1: Power Automate. Day 2: Copilot Studio agents, ending with the
assessment block: WA 1 hr + PP 1 hr, 4:00 - 6:00pm.
The Slides column maps every session to courseware/Business Process Automation with Power Automate and Copilot Studio Agents-v3.pptx
(86 slides after the 2-day restructure).

Writes: courseware/LP-<course>.docx
"""
import os, sys
SKILL = "/Users/alfredang/.claude/skills/tertiary-lesson-plan"
sys.path.insert(0, SKILL)
import prodoc
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# script lives at .claude/skills/wsq-lesson-plan/ — repo root is 3 levels up
REPO = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", ".."))
TITLE = "Business Process Automation with Power Automate and Copilot Studio Agents"
VERSION = "3.0"
COURSE_CODE = "TGS-2022017524"
ORG = "Tertiary Infotech Academy Pte Ltd"
UEN = "201200696W"
LMS = "https://lms-tms.tertiaryinfotech.com/"
VERSIONS = [
    ["1.0", "24 Jun 2026", "Initial release — 3-day lesson plan (9:00am-5:00pm).",
     "Course Development Team"],
    ["2.0", "2 Jul 2026", "WSQ revision — new course title, 9:00am-6:00pm schedule, "
     "Day-3 assessment block (WA 1 hr + PP 1 hr), slide mapping for the 108-slide deck.",
     "Course Development Team"],
    ["3.0", "3 Jul 2026", "Course restructured from 3 days to 2 days — Day 1: Power Automate, "
     "Day 2: Copilot Studio agents ending with the assessment block (WA 1 hr + PP 1 hr, 4:00-6:00pm). "
     "Modules 4-5 and Labs 12-16 retired; slide mapping updated for the 86-slide v3 deck.",
     "Course Development Team"],
]

BRAND = RGBColor(0x1F,0x6F,0xEB); DARK = RGBColor(0x16,0x1B,0x26); GREY = RGBColor(0x55,0x5B,0x66)
CENTER = WD_ALIGN_PARAGRAPH.CENTER

def _cline(doc, text, size, bold=False, color=DARK, before=0, after=4):
    p = doc.add_paragraph(); p.alignment = CENTER
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = "Arial"
    return p

def add_cover(doc):
    for _ in range(3): doc.add_paragraph()
    _cline(doc, ORG, 13, bold=True, after=2)
    _cline(doc, f"UEN: {UEN}", 10, color=GREY, after=20)
    _cline(doc, "LESSON PLAN", 26, bold=True, color=BRAND, after=12)
    _cline(doc, "For", 12, color=GREY, after=8)
    _cline(doc, TITLE, 20, bold=True, color=DARK, after=14)
    _cline(doc, f"TGS Ref No: {COURSE_CODE}", 12, color=GREY, after=6)
    _cline(doc, "Duration: 2 Days  ·  9:00am – 6:00pm", 12, color=GREY, after=20)
    _cline(doc, "Conducted by", 12, color=GREY, after=6)
    _cline(doc, ORG, 13, bold=True, after=2)
    _cline(doc, f"UEN: {UEN}", 10, color=GREY, after=18)
    _cline(doc, f"Version {VERSION}", 12, bold=True, color=BRAND)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_footer(doc):
    sec = doc.sections[0]; f = sec.footer; f.is_linked_to_previous = False
    p = f.paragraphs[0]; p.alignment = CENTER; p.text = ""
    r = p.add_run("Page "); r.font.size = Pt(9); r.font.color.rgb = GREY
    prodoc._field(p, "PAGE", "1").font.size = Pt(9)
    r2 = p.add_run(" of "); r2.font.size = Pt(9); r2.font.color.rgb = GREY
    prodoc._field(p, "NUMPAGES", "1").font.size = Pt(9)
    sp = f.add_paragraph(); sp.alignment = CENTER
    sr = sp.add_run(f"{TITLE} — Lesson Plan  ·  © {ORG}")
    sr.font.size = Pt(7.5); sr.font.color.rgb = GREY

def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); tcPr.append(shd)

def _borders(t, color="7A8190", sz=6):
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
    b = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),"single"); e.set(qn("w:sz"),str(sz))
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),color)
        b.append(e)
    tblPr.append(b)

def heading(doc, text, lvl=1):
    doc.add_paragraph(style=f"Heading {lvl}").add_run(text)
def para(doc, text):
    p = doc.add_paragraph(); p.add_run(text); return p
def bullets(doc, items):
    for it in items: doc.add_paragraph(style="List Bullet").add_run(it)

def info_table(doc, rows):
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = ""; rr = c[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(10); _shade(c[0], "EAF1FF")
        c[1].text = ""; c[1].paragraphs[0].add_run(v).font.size = Pt(10)
    _borders(t)

def schedule_table(doc, rows):
    """rows: (time, duration, activity, method, kind, slides); kind in {'topic','lab','break','assess',''}"""
    t = doc.add_table(rows=0, cols=5); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.add_row().cells
    for i, h in enumerate(["Time", "Duration", "Topic / Activity", "Method", "Slides"]):
        hdr[i].text = ""; rr = hdr[i].paragraphs[0].add_run(h)
        rr.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); _shade(hdr[i], "1F6FEB")
    for time, dur, act, method, kind, slides in rows:
        cells = t.add_row().cells
        vals = [time, dur, act, method, slides]
        tint = {"break":"FFF4E5", "topic":"EAF1FF", "lab":"E8F7EE", "assess":"FDE8E8"}.get(kind)
        for i, v in enumerate(vals):
            cells[i].text = ""; pp = cells[i].paragraphs[0]
            rr = pp.add_run(v); rr.font.size = Pt(9.5)
            if i == 2 and kind in ("topic","lab","assess"): rr.bold = True
            if tint: _shade(cells[i], tint)
    for row in t.rows:
        row.cells[0].width = Pt(76); row.cells[1].width = Pt(50)
        row.cells[3].width = Pt(88); row.cells[4].width = Pt(56)
    _borders(t)

# ---------------------------------------------------------------- schedules
LEC = "Lecture & demo"; HND = "Hands-on lab"; DISC = "Facilitated discussion"
BRK = "—"; ASMT = "Individual assessment"

DAY1 = [
    ("9:00 – 9:30",  "30 min", "Welcome, WSQ admin & digital attendance, introductions, ground rules, course overview", DISC, "topic", "1–14"),
    ("9:30 – 10:00", "30 min", "Lab 0: Environment Setup — accounts, Course Sandbox environment check", HND, "lab", "15"),
    ("10:00 – 10:45","45 min", "Module 1: Introduction to Workflow Automation — triggers, actions, outputs, steps", LEC, "topic", "16–27"),
    ("10:45 – 11:00","15 min", "Tea break", BRK, "break", "—"),
    ("11:00 – 11:30","30 min", "Module 2: Introduction to Power Automate — flow types, common triggers, connectors", LEC, "topic", "28–40"),
    ("11:30 – 12:30","60 min", "Lab 1: Automated Email Workflow (manual trigger + input → Send an email)", HND, "lab", "41–42"),
    ("12:30 – 1:30", "60 min", "Lunch", BRK, "break", "—"),
    ("1:30 – 2:30",  "60 min", "Lab 2: Excel Data Logging Workflow — build the table & flow", HND, "lab", "43"),
    ("2:30 – 3:00",  "30 min", "Lab 2 (cont.): map columns, fx date expression, test & verify rows", HND, "lab", "43"),
    ("3:00 – 3:15",  "15 min", "Tea break", BRK, "break", "—"),
    ("3:15 – 4:15",  "60 min", "Lab 3: Simple Approval Workflow — approvals & conditions", HND, "lab", "44"),
    ("4:15 – 5:15",  "60 min", "Lab 4: Scheduled Trigger Workflow — Recurrence & reminders", HND, "lab", "45"),
    ("5:15 – 6:00",  "45 min", "Lab 5: Form Submission Workflow + Day 1 recap & Q&A", HND, "lab", "46–47"),
]
DAY2 = [
    ("9:00 – 9:15",  "15 min", "Day 1 recap & Q&A", DISC, "topic", "47–49"),
    ("9:15 – 10:00", "45 min", "Module 3: Building Business Agents with Copilot Studio — prompt design for structured output", LEC, "topic", "50–59"),
    ("10:00 – 10:40","40 min", "Lab 6: Create Your First Agent — instructions, knowledge, Test pane", HND, "lab", "65–66"),
    ("10:40 – 10:55","15 min", "Tea break", BRK, "break", "—"),
    ("10:55 – 11:35","40 min", "Lab 7: Add Knowledge to Your Agent (Grounding / RAG)", HND, "lab", "67"),
    ("11:35 – 12:15","40 min", "Lab 8: Add Tools and Actions — connectors & flows as tools", HND, "lab", "68"),
    ("12:15 – 1:15", "60 min", "Lunch", BRK, "break", "—"),
    ("1:15 – 2:00",  "45 min", "Lab 9: Sales Enquiry Assistant — capture structured data", HND, "lab", "69"),
    ("2:00 – 2:50",  "50 min", "Lab 10: Procurement Request Workflow — agent calls a flow, log + notify, test end-to-end", HND, "lab", "60–64, 70"),
    ("2:50 – 3:00",  "10 min", "Tea break", BRK, "break", "—"),
    ("3:00 – 3:30",  "30 min", "Lab 11: Automated Response Generation (AI prompts)", HND, "lab", "71–72"),
    ("3:30 – 4:00",  "30 min", "Testing & debugging agents, Day 2 recap, course wrap-up & briefing for assessment", DISC, "topic", "72–85"),
    ("4:00 – 5:00",  "60 min", "Written Assessment (WA / SAQ) — open book", ASMT, "assess", "—"),
    ("5:00 – 6:00",  "60 min", "Practical Performance (PP) Assessment — open book", ASMT, "assess", "—"),
]

def total(rows):
    return sum(int(r[1].split()[0]) for r in rows)
for nm, rows in [("Day 1",DAY1),("Day 2",DAY2)]:
    assert total(rows) == 540, f"{nm} = {total(rows)} min (expected 540)"

# ---------------------------------------------------------------- build doc
doc = Document()
nrm = doc.styles["Normal"]; nrm.font.name = "Arial"; nrm.font.size = Pt(11)
prodoc.style_headings(doc)
add_cover(doc)
prodoc.add_version_control(doc, VERSIONS)
prodoc.add_toc(doc, levels="1-1")

heading(doc, "Course Overview", 1)
para(doc, "This 2-day, hands-on WSQ course teaches participants to automate real business processes by combining "
          "Microsoft Power Automate flows with AI agents built in Microsoft Copilot Studio. Day 1 covers workflow "
          "automation concepts and Power Automate; Day 2 covers business agents in Copilot Studio and connecting "
          "them to flows. Every topic follows a concept → demonstration → hands-on lab pattern; participants "
          "finish with working automations and sit the WSQ assessment at the end of Day 2.")
info_table(doc, [
    ("Course Title", TITLE),
    ("TGS Ref No", COURSE_CODE),
    ("Duration", "2 Days (9:00am – 6:00pm), incl. 1-hour lunch and tea breaks; assessment on Day 2, 4:00 – 6:00pm"),
    ("Delivery", "Instructor-led, hands-on (physical / virtual)"),
    ("Audience", "Business and operations staff who want to automate repetitive work; no coding required"),
    ("Prerequisites", "Basic Microsoft 365 familiarity (Outlook, Excel); a Power Platform environment (see Lab 0)"),
    ("Labs", "12 step-by-step labs across 2 days (Labs 0–11), plus 3 concept modules"),
    ("Assessment", "Written Assessment (SAQ) 1 hr + Practical Performance (PP) 1 hr — Day 2, 4:00 – 6:00pm, open book"),
])

heading(doc, "Learning Outcomes", 1)
para(doc, "By the end of the course, participants will be able to:")
bullets(doc, [
    "Explain business workflow automation and the Trigger → Actions → Output model.",
    "Build Power Automate flows that send emails, log data to Excel, run approvals, and use scheduled and form triggers.",
    "Create business agents in Copilot Studio, ground them with knowledge (RAG), and give them tools.",
    "Connect agents to Power Automate flows and pass structured data between them.",
    "Combine agents and flows into complete end-to-end business automations (e.g. procurement request handling).",
])

heading(doc, "Daily Schedule", 1)
para(doc, "The Slides column maps each session to the matching slides in the facilitator deck "
          "(Business Process Automation with Power Automate and Copilot Studio Agents-v3.pptx, 86 slides) so trainers can pace delivery against the deck.")
for nm, theme, rows in [
    ("Day 1 — Foundations & Power Automate", "Workflow concepts + your first five flows", DAY1),
    ("Day 2 — Building Business Agents with Copilot Studio & Assessment", "Agents, knowledge/RAG, tools, agent + flow end-to-end, then sit the assessment", DAY2),
]:
    heading(doc, nm, 2)
    para(doc, theme + ".")
    schedule_table(doc, rows)
    doc.add_paragraph()

heading(doc, "Tools & Resources", 1)
info_table(doc, [
    ("Microsoft Power Automate", "make.powerautomate.com — build and run flows"),
    ("Microsoft Copilot Studio", "copilotstudio.microsoft.com — build and publish agents"),
    ("Power Platform admin center", "admin.powerplatform.microsoft.com — create the Course Sandbox environment (Dataverse)"),
    ("Microsoft 365", "Outlook, Excel (OneDrive), Microsoft Forms"),
    ("Learner Guide", "Full step-by-step guide for every lab (LMS + course GitHub repository)"),
    ("LMS", LMS + " — courseware download, TRAQOM, assessment submission"),
])

heading(doc, "Assessment", 1)
para(doc, "The summative WSQ assessment is taken on Day 2, 4:00 – 6:00pm, and is open book "
          "(slides, Learner Guide and approved materials only):")
bullets(doc, [
    "Written Assessment (WA / SAQ) — 1 hour, 4:00 – 5:00pm. Open-ended short-answer questions testing the knowledge from Modules 1–3.",
    "Practical Performance (PP) — 1 hour, 5:00 – 6:00pm. Candidates build a working flow and agent, mirroring the hands-on labs.",
    "Assessment flow: TRAQOM survey (QR on the LMS) → Assessment digital attendance → WA then PP → submit answers on the LMS → sign the Assessment Summary Record.",
    "WSQ funding requires a minimum of 75% attendance AND a Competent (C) assessment outcome.",
    f"Courseware and the assessment are on the LMS: {LMS}",
])
para(doc, "Formative assessment is continuous: each lab includes a Checkpoint that the facilitator verifies before "
          "the class moves on, and Lab 10 (Procurement Request Workflow) serves as the integrative check that a "
          "participant can wire an agent to a flow end-to-end ahead of the Practical Performance assessment.")

add_footer(doc)
prodoc.enable_update_fields(doc)
out = os.path.join(REPO, f"courseware/LP-{TITLE}.docx")
doc.save(out)
print("Day totals:", {n: total(r) for n,r in [("D1",DAY1),("D2",DAY2)]})
print("Wrote", out)
